"""Document processing utilities for text extraction and chunking"""

import os
from typing import List, Tuple, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processor for extracting text from various document formats"""
    
    CHUNK_SIZE = 512
    CHUNK_OVERLAP = 50
    MIN_CHUNK_SIZE = 50
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text from PDF
            
        Raises:
            Exception: If PDF extraction fails
        """
        text = ""
        
        # Try PyMuPDF (fitz) first as it's generally better
        try:
            import fitz
            with fitz.open(file_path) as pdf:
                for page_num in range(len(pdf)):
                    page = pdf[page_num]
                    text += page.get_text()
                    text += f"\n[Page {page_num + 1}]\n"
            
            if text.strip():
                return text.strip()
        except ImportError:
            logger.warning("PyMuPDF not installed, falling back to pypdf")
        except Exception as e:
            logger.warning(f"PyMuPDF failed for {file_path}: {str(e)}, trying pypdf")

        # Fallback to pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                    text += f"\n[Page {i + 1}]\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """
        Extract text from text file
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read().strip()
    
    @staticmethod
    def extract_text_from_markdown(file_path: str) -> str:
        """
        Extract text from markdown file
        
        Args:
            file_path: Path to markdown file
            
        Returns:
            Markdown content
        """
        return DocumentProcessor.extract_text_from_txt(file_path)
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text from DOCX
            
        Raises:
            ImportError: If python-docx not installed
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required for DOCX processing. Install: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_doc(file_path: str) -> str:
        """
        Extract text from legacy DOC file.
        Note: python-docx only supports DOCX. 
        Legacy DOC files require external conversion or specialized libraries.
        
        Args:
            file_path: Path to DOC file
            
        Returns:
            Extracted text from DOC
        """
        try:
            from docx import Document
            # Try as docx first (in case it's just misnamed)
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception:
            logger.warning(f"Legacy DOC format detected for {file_path}. python-docx cannot process binary DOC files.")
            raise Exception(
                "Legacy .doc format is not directly supported. "
                "Please convert the file to .docx or .txt before uploading."
            )
    
    @staticmethod
    def extract_text_from_json(file_path: str) -> str:
        """
        Extract text from JSON file
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            JSON content as text
        """
        import json
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, dict):
                    return json.dumps(data, indent=2)
                elif isinstance(data, list):
                    return "\n".join([json.dumps(item, indent=2) for item in data])
                else:
                    return str(data)
        except Exception as e:
            logger.error(f"JSON parsing failed for {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from JSON: {str(e)}")
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from document based on file extension
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If file type not supported
            Exception: If extraction fails
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        extractors = {
            '.pdf': DocumentProcessor.extract_text_from_pdf,
            '.txt': DocumentProcessor.extract_text_from_txt,
            '.md': DocumentProcessor.extract_text_from_markdown,
            '.docx': DocumentProcessor.extract_text_from_docx,
            '.doc': DocumentProcessor.extract_text_from_doc,
            '.json': DocumentProcessor.extract_text_from_json,
        }
        
        if extension not in extractors:
            raise ValueError(f"Unsupported file type: {extension}")
        
        logger.info(f"Extracting text from {path.name} ({extension})")
        return extractors[extension](file_path)


class ChunkingStrategy:
    """Text chunking strategies"""
    
    @staticmethod
    def chunk_by_size(
        text: str,
        chunk_size: int = 512,
        overlap: int = 50
    ) -> List[str]:
        """
        Chunk text by character size with overlap
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk in characters
            overlap: Overlap between chunks in characters
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end]
            
            if len(chunk) > DocumentProcessor.MIN_CHUNK_SIZE:
                chunks.append(chunk)
            
            start += chunk_size - overlap
        
        return chunks
    
    @staticmethod
    def chunk_by_tokens(
        text: str,
        max_tokens: int = 512,
        overlap_tokens: int = 50
    ) -> List[str]:
        """
        Chunk text by token count (approximate)
        
        Args:
            text: Text to chunk
            max_tokens: Maximum tokens per chunk (approximate)
            overlap_tokens: Overlap in tokens
            
        Returns:
            List of text chunks
        """
        try:
            import tiktoken
            encoding = tiktoken.get_encoding("cl100k_base")
        except ImportError:
            logger.warning("tiktoken not available, using character-based estimation")
            avg_token_length = 4
            char_size = max_tokens * avg_token_length
            overlap_chars = overlap_tokens * avg_token_length
            return ChunkingStrategy.chunk_by_size(text, char_size, overlap_chars)
        
        chunks = []
        tokens = encoding.encode(text)
        
        start = 0
        while start < len(tokens):
            end = min(start + max_tokens, len(tokens))
            token_chunk = tokens[start:end]
            chunk_text = encoding.decode(token_chunk)
            
            if len(chunk_text) > DocumentProcessor.MIN_CHUNK_SIZE:
                chunks.append(chunk_text)
            
            start += max_tokens - overlap_tokens
        
        return chunks
    
    @staticmethod
    def chunk_by_sentences(
        text: str,
        sentences_per_chunk: int = 3,
        overlap_sentences: int = 1
    ) -> List[str]:
        """
        Chunk text by sentences
        
        Args:
            text: Text to chunk
            sentences_per_chunk: Number of sentences per chunk
            overlap_sentences: Overlap in sentences
            
        Returns:
            List of text chunks
        """
        try:
            import nltk
            nltk.download('punkt', quiet=True)
            from nltk.tokenize import sent_tokenize
        except ImportError:
            logger.warning("nltk not available, using period-based sentence splitting")
            sentences = text.split('. ')
        else:
            sentences = sent_tokenize(text)
        
        chunks = []
        start = 0
        
        while start < len(sentences):
            end = min(start + sentences_per_chunk, len(sentences))
            chunk_sentences = sentences[start:end]
            chunk = '. '.join(chunk_sentences)
            
            if len(chunk) > DocumentProcessor.MIN_CHUNK_SIZE:
                chunks.append(chunk)
            
            start += sentences_per_chunk - overlap_sentences
        
        return chunks
    
    @staticmethod
    def chunk_by_paragraphs(text: str) -> List[str]:
        """
        Chunk text by paragraphs
        
        Args:
            text: Text to chunk
            
        Returns:
            List of paragraphs as chunks
        """
        chunks = [p.strip() for p in text.split('\n\n') if p.strip()]
        return [c for c in chunks if len(c) > DocumentProcessor.MIN_CHUNK_SIZE]


class TokenCounter:
    """Token counting utilities"""
    
    @staticmethod
    def count_tokens(text: str) -> int:
        """
        Count approximate tokens in text
        
        Args:
            text: Text to count
            
        Returns:
            Estimated token count
        """
        try:
            import tiktoken
            encoding = tiktoken.get_encoding("cl100k_base")
            return len(encoding.encode(text))
        except ImportError:
            logger.warning("tiktoken not available, using character-based estimation")
            return len(text) // 4
    
    @staticmethod
    def count_words(text: str) -> int:
        """
        Count words in text
        
        Args:
            text: Text to count
            
        Returns:
            Word count
        """
        return len(text.split())
